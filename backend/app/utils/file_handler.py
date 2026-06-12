"""File handling utilities"""

import fitz  # PyMuPDF
from docx import Document
import os


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        pdf_document = fitz.open(file_path)
        for page in pdf_document:
            text += page.get_text()
        pdf_document.close()
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")
    return text


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")
    return text


def extract_text_from_file(file_path: str, file_type: str) -> str:
    """Extract text from file based on type"""
    file_type = file_type.lower()
    
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type in ["docx", "doc"]:
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def clean_text(text: str) -> str:
    """Clean and normalize text"""
    # Remove extra whitespace
    text = " ".join(text.split())
    return text


def extract_keywords(text: str, num_keywords: int = 20) -> list:
    """Extract keywords from text"""
    # Simple keyword extraction based on word frequency
    words = text.lower().split()
    
    # Remove common words
    stop_words = {
        "the", "a", "an", "and", "or", "but", "in", "on", "at", "to", "for",
        "of", "with", "by", "from", "up", "about", "into", "through", "during",
        "is", "are", "was", "were", "be", "been", "being", "have", "has", "had",
        "do", "does", "did", "will", "would", "could", "should", "may", "might",
        "must", "can", "this", "that", "these", "those", "i", "you", "he", "she",
        "it", "we", "they", "what", "which", "who", "when", "where", "why", "how"
    }
    
    filtered_words = [w for w in words if w not in stop_words and len(w) > 3]
    
    # Count word frequency
    word_freq = {}
    for word in filtered_words:
        word_freq[word] = word_freq.get(word, 0) + 1
    
    # Sort by frequency
    sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
    
    # Return top keywords
    return [word[0] for word in sorted_words[:num_keywords]]
